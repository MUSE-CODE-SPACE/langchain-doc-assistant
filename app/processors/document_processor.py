"""
Document Processor - Handles various document formats.

Provides a unified interface for parsing PDF, DOCX, TXT, Markdown, CSV (and
HTML/XLSX best-effort) inputs into a ``ProcessedDocument`` carrying raw text,
chunks and lightweight metadata. The processors degrade gracefully when their
optional native libraries are missing.
"""

from __future__ import annotations

import contextlib
import csv
import hashlib
import io
import logging
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import UTC, datetime
from enum import Enum
from typing import Any, BinaryIO, ClassVar

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "markdown"
    HTML = "html"
    CSV = "csv"
    XLSX = "xlsx"
    JSON = "json"
    UNKNOWN = "unknown"


@dataclass
class DocumentMetadata:
    """Document metadata."""
    filename: str
    doc_type: DocumentType
    size_bytes: int
    page_count: int = 1
    word_count: int = 0
    char_count: int = 0
    created_at: datetime = field(default_factory=lambda: datetime.now(UTC))
    author: str | None = None
    title: str | None = None
    language: str = "en"
    checksum: str = ""
    custom_metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentChunk:
    """A chunk of document content."""
    content: str
    chunk_id: str
    document_id: str
    page_number: int = 1
    chunk_index: int = 0
    start_char: int = 0
    end_char: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)
    embedding: list[float] | None = None

    def to_dict(self) -> dict[str, Any]:
        return {
            "content": self.content,
            "chunk_id": self.chunk_id,
            "document_id": self.document_id,
            "page_number": self.page_number,
            "chunk_index": self.chunk_index,
            "metadata": self.metadata,
        }


@dataclass
class ProcessedDocument:
    """Fully processed document."""
    document_id: str
    metadata: DocumentMetadata
    raw_content: str
    chunks: list[DocumentChunk]
    tables: list[dict[str, Any]] = field(default_factory=list)
    images: list[dict[str, Any]] = field(default_factory=list)
    sections: list[dict[str, Any]] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _as_bytes(content: str | bytes | BinaryIO) -> bytes:
    """Best-effort conversion of mixed content inputs to bytes."""
    if isinstance(content, bytes):
        return content
    if isinstance(content, str):
        return content.encode("utf-8", errors="ignore")
    if hasattr(content, "read"):
        data = content.read()
        with contextlib.suppress(Exception):
            content.seek(0)
        if isinstance(data, str):
            return data.encode("utf-8", errors="ignore")
        return data
    return b""


# ---------------------------------------------------------------------------
# Processors
# ---------------------------------------------------------------------------
class BaseDocumentProcessor(ABC):
    """Abstract base class for document processors."""

    @abstractmethod
    def process(
        self, content: str | bytes | BinaryIO, filename: str
    ) -> ProcessedDocument:
        """Process document content."""

    @abstractmethod
    def extract_text(self, content: str | bytes | BinaryIO) -> str:
        """Extract plain text from document."""

    def generate_document_id(self, content: bytes, filename: str) -> str:
        """Generate unique document ID."""
        hash_input = f"{filename}{len(content)}{datetime.now(UTC).isoformat()}"
        return hashlib.sha256(hash_input.encode()).hexdigest()[:16]


class TextDocumentProcessor(BaseDocumentProcessor):
    """Process plain text documents."""

    def process(
        self, content: str | bytes | BinaryIO, filename: str
    ) -> ProcessedDocument:
        text = self.extract_text(content)
        content_bytes = _as_bytes(content) or text.encode("utf-8", errors="ignore")

        doc_id = self.generate_document_id(content_bytes, filename)

        metadata = DocumentMetadata(
            filename=filename,
            doc_type=DocumentType.TXT,
            size_bytes=len(content_bytes),
            word_count=len(text.split()),
            char_count=len(text),
            checksum=hashlib.md5(text.encode("utf-8", errors="ignore")).hexdigest(),
        )

        chunks = self._create_chunks(text, doc_id)

        return ProcessedDocument(
            document_id=doc_id,
            metadata=metadata,
            raw_content=text,
            chunks=chunks,
        )

    def extract_text(self, content: str | bytes | BinaryIO) -> str:
        if isinstance(content, str):
            return content
        if isinstance(content, bytes):
            return content.decode("utf-8", errors="ignore")
        return content.read().decode("utf-8", errors="ignore")

    def _create_chunks(
        self,
        text: str,
        doc_id: str,
        chunk_size: int = 1000,
        overlap: int = 200,
    ) -> list[DocumentChunk]:
        """Split text into overlapping chunks by paragraphs."""
        chunks: list[DocumentChunk] = []
        paragraphs = text.split("\n\n")

        current_chunk = ""
        current_start = 0
        chunk_index = 0

        for para in paragraphs:
            if len(current_chunk) + len(para) < chunk_size:
                current_chunk += para + "\n\n"
                continue

            if current_chunk.strip():
                chunks.append(
                    DocumentChunk(
                        content=current_chunk.strip(),
                        chunk_id=f"{doc_id}_chunk_{chunk_index}",
                        document_id=doc_id,
                        chunk_index=chunk_index,
                        start_char=current_start,
                        end_char=current_start + len(current_chunk),
                    )
                )
                chunk_index += 1

            overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else ""
            current_start = current_start + len(current_chunk) - len(overlap_text)
            current_chunk = overlap_text + para + "\n\n"

        if current_chunk.strip():
            chunks.append(
                DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_id=f"{doc_id}_chunk_{chunk_index}",
                    document_id=doc_id,
                    chunk_index=chunk_index,
                    start_char=current_start,
                    end_char=current_start + len(current_chunk),
                )
            )

        return chunks


class MarkdownProcessor(TextDocumentProcessor):
    """Process Markdown documents."""

    def process(
        self, content: str | bytes | BinaryIO, filename: str
    ) -> ProcessedDocument:
        doc = super().process(content, filename)
        doc.metadata.doc_type = DocumentType.MD
        doc.sections = self._extract_sections(doc.raw_content)
        return doc

    def _extract_sections(self, text: str) -> list[dict[str, Any]]:
        """Extract markdown sections by headers."""
        sections: list[dict[str, Any]] = []
        current_section: dict[str, Any] = {
            "title": "Introduction",
            "level": 0,
            "content": "",
        }

        for line in text.split("\n"):
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if header_match:
                if current_section["content"].strip():
                    sections.append(current_section)
                level = len(header_match.group(1))
                title = header_match.group(2)
                current_section = {"title": title, "level": level, "content": ""}
            else:
                current_section["content"] += line + "\n"

        if current_section["content"].strip():
            sections.append(current_section)

        return sections


class PDFProcessor(BaseDocumentProcessor):
    """Process PDF documents using pypdf."""

    def process(
        self, content: str | bytes | BinaryIO, filename: str
    ) -> ProcessedDocument:
        text = self.extract_text(content)
        content_bytes = _as_bytes(content)
        doc_id = self.generate_document_id(content_bytes, filename)

        page_count = text.count("\f") + 1 if "\f" in text else max(text.count("[Page "), 1)

        metadata = DocumentMetadata(
            filename=filename,
            doc_type=DocumentType.PDF,
            size_bytes=len(content_bytes),
            page_count=page_count,
            word_count=len(text.split()),
            char_count=len(text),
            checksum=hashlib.md5(content_bytes).hexdigest(),
        )

        chunks = self._create_chunks_with_pages(text, doc_id)

        return ProcessedDocument(
            document_id=doc_id,
            metadata=metadata,
            raw_content=text,
            chunks=chunks,
        )

    def extract_text(self, content: str | bytes | BinaryIO) -> str:
        """Extract text from PDF using pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf not installed; PDF extraction unavailable.")
            return "[PDF processing requires pypdf library]"

        try:
            if isinstance(content, bytes):
                pdf_file: Any = io.BytesIO(content)
            elif isinstance(content, str):
                pdf_file = open(content, "rb")  # noqa: SIM115 - closed by GC
            else:
                pdf_file = content

            reader = PdfReader(pdf_file)
            parts: list[str] = []
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                parts.append(f"[Page {page_num}]\n{page_text}")
            return "\n\n".join(parts)
        except Exception as exc:
            logger.warning("Failed to extract PDF: %s", exc)
            return f"[Error extracting PDF text: {exc}]"

    def _create_chunks_with_pages(
        self,
        text: str,
        doc_id: str,
        chunk_size: int = 1000,
    ) -> list[DocumentChunk]:
        """Create chunks while tracking page numbers."""
        chunks: list[DocumentChunk] = []
        page_pattern = r"\[Page (\d+)\]"
        pages = re.split(page_pattern, text)

        chunk_index = 0
        for i in range(1, len(pages), 2):
            page_num = int(pages[i])
            page_content = pages[i + 1] if i + 1 < len(pages) else ""

            for j in range(0, len(page_content), chunk_size):
                chunk_text = page_content[j : j + chunk_size].strip()
                if chunk_text:
                    chunks.append(
                        DocumentChunk(
                            content=chunk_text,
                            chunk_id=f"{doc_id}_chunk_{chunk_index}",
                            document_id=doc_id,
                            page_number=page_num,
                            chunk_index=chunk_index,
                            metadata={"page": page_num},
                        )
                    )
                    chunk_index += 1

        if not chunks and text.strip():
            # Fallback to a flat chunker for PDFs without page markers.
            return TextDocumentProcessor()._create_chunks(text, doc_id)
        return chunks


class DocxProcessor(BaseDocumentProcessor):
    """Process Word documents using python-docx."""

    def process(
        self, content: str | bytes | BinaryIO, filename: str
    ) -> ProcessedDocument:
        text, tables = self._extract_content(content)
        content_bytes = _as_bytes(content)
        doc_id = self.generate_document_id(content_bytes, filename)

        metadata = DocumentMetadata(
            filename=filename,
            doc_type=DocumentType.DOCX,
            size_bytes=len(content_bytes),
            word_count=len(text.split()),
            char_count=len(text),
            checksum=hashlib.md5(content_bytes).hexdigest(),
        )

        chunks = TextDocumentProcessor()._create_chunks(text, doc_id)

        return ProcessedDocument(
            document_id=doc_id,
            metadata=metadata,
            raw_content=text,
            chunks=chunks,
            tables=tables,
        )

    def extract_text(self, content: str | bytes | BinaryIO) -> str:
        text, _ = self._extract_content(content)
        return text

    def _extract_content(
        self, content: str | bytes | BinaryIO
    ) -> tuple[str, list[dict[str, Any]]]:
        """Extract text and tables from DOCX."""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed; DOCX extraction unavailable.")
            return "[DOCX processing requires python-docx library]", []

        try:
            if isinstance(content, bytes):
                doc_file: Any = io.BytesIO(content)
            elif isinstance(content, str):
                doc_file = content
            else:
                doc_file = content

            doc = Document(doc_file)

            text_parts: list[str] = []
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    level_ch = style_name[-1]
                    level = int(level_ch) if level_ch.isdigit() else 1
                    text_parts.append(f"\n{'#' * level} {para.text}\n")
                else:
                    text_parts.append(para.text)

            tables: list[dict[str, Any]] = []
            for table in doc.tables:
                table_data = [[cell.text for cell in row.cells] for row in table.rows]
                tables.append(
                    {
                        "rows": len(table.rows),
                        "cols": len(table.columns),
                        "data": table_data,
                    }
                )

            return "\n\n".join(text_parts), tables
        except Exception as exc:
            logger.warning("Failed to extract DOCX: %s", exc)
            return f"[Error extracting DOCX: {exc}]", []


class CSVProcessor(BaseDocumentProcessor):
    """Process CSV documents."""

    def process(
        self, content: str | bytes | BinaryIO, filename: str
    ) -> ProcessedDocument:
        text = self.extract_text(content)
        tables = self._parse_csv(text)
        content_bytes = _as_bytes(content) or text.encode("utf-8", errors="ignore")
        doc_id = self.generate_document_id(content_bytes, filename)

        metadata = DocumentMetadata(
            filename=filename,
            doc_type=DocumentType.CSV,
            size_bytes=len(content_bytes),
            word_count=len(text.split()),
            char_count=len(text),
            checksum=hashlib.md5(content_bytes).hexdigest(),
        )

        chunks = self._create_row_chunks(tables, doc_id)

        return ProcessedDocument(
            document_id=doc_id,
            metadata=metadata,
            raw_content=text,
            chunks=chunks,
            tables=tables,
        )

    def extract_text(self, content: str | bytes | BinaryIO) -> str:
        if isinstance(content, str):
            return content
        if isinstance(content, bytes):
            return content.decode("utf-8", errors="ignore")
        return content.read().decode("utf-8", errors="ignore")

    def _parse_csv(self, text: str) -> list[dict[str, Any]]:
        """Parse CSV into structured table."""
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return []

        headers = rows[0]
        data = rows[1:] if len(rows) > 1 else []

        return [
            {
                "headers": headers,
                "rows": len(data),
                "cols": len(headers),
                "data": data,
            }
        ]

    def _create_row_chunks(
        self,
        tables: list[dict[str, Any]],
        doc_id: str,
        rows_per_chunk: int = 50,
    ) -> list[DocumentChunk]:
        """Create chunks from table rows."""
        chunks: list[DocumentChunk] = []

        for table in tables:
            headers = table.get("headers", [])
            data = table.get("data", [])

            for i in range(0, len(data), rows_per_chunk):
                chunk_rows = data[i : i + rows_per_chunk]

                content_parts = [" | ".join(headers), "-" * 50]
                content_parts.extend(
                    " | ".join(str(cell) for cell in row) for row in chunk_rows
                )

                chunks.append(
                    DocumentChunk(
                        content="\n".join(content_parts),
                        chunk_id=f"{doc_id}_chunk_{len(chunks)}",
                        document_id=doc_id,
                        chunk_index=len(chunks),
                        metadata={
                            "row_start": i,
                            "row_end": i + len(chunk_rows),
                            "headers": headers,
                        },
                    )
                )

        return chunks


class HTMLProcessor(TextDocumentProcessor):
    """Process HTML by stripping tags with BeautifulSoup when available."""

    def process(
        self, content: str | bytes | BinaryIO, filename: str
    ) -> ProcessedDocument:
        doc = super().process(content, filename)
        doc.metadata.doc_type = DocumentType.HTML
        return doc

    def extract_text(self, content: str | bytes | BinaryIO) -> str:
        raw = super().extract_text(content)
        try:
            from bs4 import BeautifulSoup

            return BeautifulSoup(raw, "html.parser").get_text(separator="\n").strip()
        except Exception:
            return re.sub(r"<[^>]+>", "", raw)


class XLSXProcessor(BaseDocumentProcessor):
    """Process XLSX spreadsheets via openpyxl."""

    def process(
        self, content: str | bytes | BinaryIO, filename: str
    ) -> ProcessedDocument:
        text, tables = self._extract_content(content)
        content_bytes = _as_bytes(content)
        doc_id = self.generate_document_id(content_bytes, filename)

        metadata = DocumentMetadata(
            filename=filename,
            doc_type=DocumentType.XLSX,
            size_bytes=len(content_bytes),
            word_count=len(text.split()),
            char_count=len(text),
            checksum=hashlib.md5(content_bytes).hexdigest(),
        )

        chunks = TextDocumentProcessor()._create_chunks(text, doc_id)
        return ProcessedDocument(
            document_id=doc_id,
            metadata=metadata,
            raw_content=text,
            chunks=chunks,
            tables=tables,
        )

    def extract_text(self, content: str | bytes | BinaryIO) -> str:
        text, _ = self._extract_content(content)
        return text

    def _extract_content(
        self, content: str | bytes | BinaryIO
    ) -> tuple[str, list[dict[str, Any]]]:
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.warning("openpyxl not installed; XLSX extraction unavailable.")
            return "[XLSX processing requires openpyxl library]", []

        try:
            buf = io.BytesIO(_as_bytes(content))
            wb = load_workbook(buf, data_only=True, read_only=True)
            text_parts: list[str] = []
            tables: list[dict[str, Any]] = []
            for sheet in wb.worksheets:
                text_parts.append(f"# Sheet: {sheet.title}")
                rows = [
                    [str(cell) if cell is not None else "" for cell in row]
                    for row in sheet.iter_rows(values_only=True)
                ]
                for row in rows:
                    text_parts.append(" | ".join(row))
                if rows:
                    tables.append(
                        {
                            "sheet": sheet.title,
                            "rows": len(rows),
                            "cols": len(rows[0]),
                            "data": rows,
                        }
                    )
            return "\n".join(text_parts), tables
        except Exception as exc:
            logger.warning("Failed to extract XLSX: %s", exc)
            return f"[Error extracting XLSX: {exc}]", []


class DocumentProcessorFactory:
    """Factory for creating document processors."""

    _processors: ClassVar[dict[DocumentType, type[BaseDocumentProcessor]]] = {
        DocumentType.TXT: TextDocumentProcessor,
        DocumentType.MD: MarkdownProcessor,
        DocumentType.PDF: PDFProcessor,
        DocumentType.DOCX: DocxProcessor,
        DocumentType.CSV: CSVProcessor,
        DocumentType.HTML: HTMLProcessor,
        DocumentType.XLSX: XLSXProcessor,
    }

    @classmethod
    def get_processor(cls, doc_type: DocumentType) -> BaseDocumentProcessor:
        """Get processor for document type."""
        processor_class = cls._processors.get(doc_type, TextDocumentProcessor)
        return processor_class()

    @classmethod
    def detect_type(cls, filename: str) -> DocumentType:
        """Detect document type from filename."""
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        type_map = {
            "pdf": DocumentType.PDF,
            "docx": DocumentType.DOCX,
            "doc": DocumentType.DOCX,
            "txt": DocumentType.TXT,
            "md": DocumentType.MD,
            "markdown": DocumentType.MD,
            "csv": DocumentType.CSV,
            "xlsx": DocumentType.XLSX,
            "html": DocumentType.HTML,
            "htm": DocumentType.HTML,
            "json": DocumentType.JSON,
        }
        return type_map.get(ext, DocumentType.UNKNOWN)

    @classmethod
    def process_document(
        cls,
        content: str | bytes | BinaryIO,
        filename: str,
    ) -> ProcessedDocument:
        """Process document with automatic type detection."""
        doc_type = cls.detect_type(filename)
        processor = cls.get_processor(doc_type)
        return processor.process(content, filename)
